"""DOCX file processor using python-docx."""

import docx
import docx.opc.exceptions

from file_processors.base_processor import BaseFileProcessor


class DocxProcessor(BaseFileProcessor):
    """Processor for DOCX files.

    Extracts text from paragraphs, tables, and section headers/footers.

    Paragraphs are joined with newlines (rather than concatenated directly) so that
    entities never fuse across paragraph boundaries — e.g. a name ending one paragraph
    and an address starting the next must not become a single token that NLP models
    misread.  Table cells are emitted row by row with tab separators so that the
    header→value relationship survives for downstream context-aware detection.
    """

    def extract_text(self, file_path: str) -> str:
        """Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content as a string

        Raises:
            docx.opc.exceptions.PackageNotFoundError: If DOCX is empty or protected
            PermissionError: If file cannot be accessed
            FileNotFoundError: If file does not exist
            Exception: For other DOCX processing errors
        """
        doc: docx.Document = docx.Document(file_path)
        parts: list[str] = []

        # Body paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        # Body tables (cell text preserves the header→value relationship per row)
        for table in doc.tables:
            parts.extend(self._extract_table(table))

        # Section headers and footers (deduplicated: sections often link to the
        # previous one, repeating the same header/footer text on every page).
        seen_hf: set[str] = set()
        for section in doc.sections:
            for container in (section.header, section.footer):
                for paragraph in container.paragraphs:
                    text = paragraph.text.strip()
                    if text and text not in seen_hf:
                        seen_hf.add(text)
                        parts.append(text)

        return "\n".join(parts)

    @staticmethod
    def _extract_table(table) -> list[str]:
        """Return one tab-joined string per table row with non-empty cells."""
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        return rows

    @staticmethod
    def can_process(extension: str) -> bool:
        """Check if this processor can handle DOCX files."""
        return extension.lower() == ".docx"
